import json
import boto3
import os
import time
import urllib.request
import re
import subprocess
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# ================= CONFIG =================
BOT_TOKEN = os.environ.get("BOT_TOKEN")
if not BOT_TOKEN:
    raise RuntimeError("BOT_TOKEN environment variable not set")

TABLE_NAME = "quotation_bot_sessions"
TEMPLATE_BUCKET = "utl-quote-template-bucket"
OUTPUT_BUCKET = "quote-output-bucket"
TEMPLATE_KEY = "template.docx"

dynamodb = boto3.resource("dynamodb")
table = dynamodb.Table(TABLE_NAME)
s3 = boto3.client("s3")

# ================= FLOW ORDER =================
FLOW_STEPS = [
    "CLIENT_NAME",
    "CAPACITY",
    "SANCTIONED_LOAD",
    "SOLAR_PANEL_MODEL",
    "SPV_MODULE",
    "INVERTER",
    "INVERTER_TYPE",
    "NO_INVERTER",
    "PHASE",
    "NO_PANELS",
    "PRICE"
]

# ================= OPTIONS =================
OPTIONS = {
    "CAPACITY": ["3 KW", "5 KW", "6 KW", "10 KW"],
    "SOLAR_PANEL_MODEL": [
        "UTL 580 Watt TOPCon Bifacial",
        "575 Watt TOPCon DCR Bi-Facial Dual Glass",
        "590 Watt N-Type TOPCon Solar Module",
        "530 Watt Mono PERC Solar Panel"
    ],
    "SPV_MODULE": [
        "Mono Half Cut",
        "Mono",
        "TOPCon Bi-Facial",
        "Mono PERC",
        "TOPCon DCR Solar"
    ],
    "INVERTER_TYPE": ["On-Grid", "Hybrid"],
    "PHASE": ["Single Phase", "Three Phase"]
}

# ================= TELEGRAM =================
def tg_send(chat_id, text, keyboard=None):
    payload = {"chat_id": chat_id, "text": text}
    if keyboard:
        payload["reply_markup"] = keyboard

    req = urllib.request.Request(
        f"https://api.telegram.org/bot{BOT_TOKEN}/sendMessage",
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"}
    )
    urllib.request.urlopen(req)

def tg_send_html(chat_id, html_text):
    payload = {
        "chat_id": chat_id,
        "text": html_text,
        "parse_mode": "HTML"
    }
    req = urllib.request.Request(
        f"https://api.telegram.org/bot{BOT_TOKEN}/sendMessage",
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"}
    )
    urllib.request.urlopen(req)

def build_keyboard(field, items, cols=2):
    keyboard, row = [], []
    for i, item in enumerate(items):
        row.append({"text": item, "callback_data": f"{field}__{i}"})
        if len(row) == cols:
            keyboard.append(row)
            row = []
    if row:
        keyboard.append(row)
    keyboard.append([{"text": "Other", "callback_data": f"{field}__OTHER"}])
    return {"inline_keyboard": keyboard}

# ================= SESSION =================
def get_session(chat_id):
    return table.get_item(Key={"chat_id": str(chat_id)}).get("Item")

def save_session(chat_id, data):
    data["chat_id"] = str(chat_id)
    data["ttl"] = int(time.time()) + 86400
    table.put_item(Item=data)

def clear_session(chat_id):
    table.delete_item(Key={"chat_id": str(chat_id)})

def next_step(step):
    idx = FLOW_STEPS.index(step)
    return FLOW_STEPS[idx + 1] if idx + 1 < len(FLOW_STEPS) else None

# ================= PRICE TO WORDS =================
def number_to_words(n):
    ones = ["","One","Two","Three","Four","Five","Six","Seven","Eight","Nine",
            "Ten","Eleven","Twelve","Thirteen","Fourteen","Fifteen",
            "Sixteen","Seventeen","Eighteen","Nineteen"]
    tens = ["","","Twenty","Thirty","Forty","Fifty","Sixty","Seventy","Eighty","Ninety"]

    def two(x): return ones[x] if x < 20 else tens[x//10] + (" " + ones[x%10] if x%10 else "")
    def three(x): return two(x) if x < 100 else ones[x//100] + " Hundred " + two(x%100)

    res = ""
    for div, name in [(10000000,"Crore"),(100000,"Lakh"),(1000,"Thousand")]:
        if n >= div:
            res += three(n//div) + f" {name} "
            n %= div
    if n > 0:
        res += three(n)
    return res.strip() + " Only"

# ================= DOCX REPLACEMENT =================
def replace_docx(doc, data):
    data = {k: str(v) for k, v in data.items()}

    def replace_in_paragraph(p):
        full_text = "".join(run.text for run in p.runs)
        updated = False

        for k, v in data.items():
            placeholder = f"{{{{{k}}}}}"
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, v)
                updated = True

        if not updated:
            return

        for run in p.runs:
            run.text = ""

        run = p.runs[0]
        run.text = full_text
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    for p in doc.paragraphs:
        if p.runs:
            replace_in_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.runs:
                        replace_in_paragraph(p)

# ================= DOCX → PDF =================
def convert_docx_to_pdf(docx_path):
    subprocess.run(
        [
            "soffice",
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to", "pdf",
            docx_path,
            "--outdir", "/tmp"
        ],
        check=True
    )

# ================= ASK NEXT =================
def ask_next(chat_id, session):
    step = session["step"]
    if step in OPTIONS:
        tg_send(chat_id, f"Select {step.replace('_',' ').title()}:", build_keyboard(step, OPTIONS[step]))
    elif step in ["SANCTIONED_LOAD", "INVERTER", "NO_INVERTER", "NO_PANELS"]:
        tg_send(chat_id, f"Enter {step.replace('_',' ').title()}:")
    elif step == "PRICE":
        tg_send(chat_id, "Enter total price (numbers only):")
    elif step is None:
        generate_files(chat_id, session)

# ================= FINAL FILE GENERATION =================
def generate_files(chat_id, session):
    session["PRICE_IN_WORDS"] = number_to_words(int(session["PRICE"]))
    session["DATE"] = datetime.now().strftime("%d-%m-%Y")

    s3.download_file(TEMPLATE_BUCKET, TEMPLATE_KEY, "/tmp/template.docx")
    doc = Document("/tmp/template.docx")
    replace_docx(doc, session)

    base = f"QUOTE_{session['CLIENT_NAME'].replace(' ','_')}_{session['CAPACITY'].replace(' ','_')}"
    docx_path = f"/tmp/{base}.docx"
    pdf_path = f"/tmp/{base}.pdf"

    doc.save(docx_path)
    s3.upload_file(docx_path, OUTPUT_BUCKET, f"{base}.docx")

    convert_docx_to_pdf(docx_path)
    s3.upload_file(pdf_path, OUTPUT_BUCKET, f"{base}.pdf")

    docx_url = s3.generate_presigned_url("get_object",
        Params={"Bucket": OUTPUT_BUCKET, "Key": f"{base}.docx"}, ExpiresIn=900)
    pdf_url = s3.generate_presigned_url("get_object",
        Params={"Bucket": OUTPUT_BUCKET, "Key": f"{base}.pdf"}, ExpiresIn=900)

    tg_send_html(chat_id, f"""
<b>Quotation Ready 📄</b>

📎 <a href="{docx_url}">Download DOCX</a>
📑 <a href="{pdf_url}">Download PDF</a>
""")

    clear_session(chat_id)

# ================= MAIN HANDLER =================
def lambda_handler(event, context):
    body = json.loads(event.get("body", "{}"))

    if "callback_query" in body:
        cq = body["callback_query"]
        chat_id = cq["message"]["chat"]["id"]
        field, value = cq["data"].split("__")
        session = get_session(chat_id) or {}

        if value == "OTHER":
            session["step"] = f"OTHER_{field}"
            save_session(chat_id, session)
            tg_send(chat_id, f"Please enter {field.replace('_',' ').title()}:")
            return {"statusCode": 200}

        session[field] = OPTIONS[field][int(value)]
        session["step"] = next_step(field)
        save_session(chat_id, session)
        ask_next(chat_id, session)
        return {"statusCode": 200}

    message = body.get("message")
    if not message:
        return {"statusCode": 200}

    chat_id = message["chat"]["id"]
    text = message.get("text","").strip()

    if text == "/quote":
        save_session(chat_id, {"step": "CLIENT_NAME"})
        tg_send(chat_id, "Enter customer name:")
        return {"statusCode": 200}

    session = get_session(chat_id)
    if not session:
        tg_send(chat_id, "Send /quote to start quotation")
        return {"statusCode": 200}

    step = session["step"]

    if step.startswith("OTHER_"):
        field = step.replace("OTHER_", "")
        session[field] = text
        session["step"] = next_step(field)
        save_session(chat_id, session)
        ask_next(chat_id, session)
        return {"statusCode": 200}

    if step in ["CLIENT_NAME","SANCTIONED_LOAD","INVERTER","NO_INVERTER","NO_PANELS"]:
        session[step] = text
        session["step"] = next_step(step)
        save_session(chat_id, session)
        ask_next(chat_id, session)
        return {"statusCode": 200}

    if step == "PRICE":
        clean = re.sub(r"[^\d]", "", text)
        if not clean:
            tg_send(chat_id, "❌ Invalid price. Example: 350000")
            return {"statusCode": 200}

        session["PRICE"] = clean
        save_session(chat_id, session)
        generate_files(chat_id, session)
        return {"statusCode": 200}

    return {"statusCode": 200}
